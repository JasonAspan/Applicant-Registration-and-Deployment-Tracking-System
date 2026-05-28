from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("ARTS Deployment Procedure.docx")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(46, 116, 181 if level <= 2 else 120)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.add_run(item)


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = Inches(widths[i])
            set_cell_text(cells[i], value)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.25


def build():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.header_distance = Inches(0.4)
    section.footer_distance = Inches(0.4)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    r = title.add_run("ARTS Deployment Procedure")
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("Ubuntu server, Hostinger domain, Cloudflare, Nginx, Flask/Gunicorn, PostgreSQL, and Synology backup/storage")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(85, 85, 85)

    add_heading(doc, "1. Target Architecture", 1)
    add_bullets(doc, [
        "Users access https://yourdomain.com through Cloudflare.",
        "Cloudflare proxies traffic to your public static IP and provides DNS, SSL edge protection, DDoS protection, and security rules.",
        "Your router forwards only TCP 80 and 443 to the Ubuntu server.",
        "Nginx receives public traffic and reverse-proxies requests to Gunicorn on 127.0.0.1:5000.",
        "Gunicorn runs the Flask ARTS application from app:app.",
        "PostgreSQL stores application data. The current system stores applicant PDFs and profile images inside PostgreSQL binary columns.",
        "Synology receives scheduled database backups. Optional future app changes can store uploaded PDFs directly on a Synology-mounted share.",
    ])

    add_heading(doc, "2. Things To Buy Or Prepare", 1)
    add_table(
        doc,
        ["Item", "Purpose"],
        [
            ("Domain from Hostinger", "Public name for the ARTS system."),
            ("Static public IPv4 from ISP", "Required for reliable self-hosting from your custom Ubuntu PC."),
            ("Cloudflare Free or Pro", "DNS, CDN, DDoS protection, SSL, WAF/rate controls. Start Free; consider Pro for heavier traffic."),
            ("UPS battery backup", "Keeps Ubuntu server, router, modem, and Synology online during short power loss."),
            ("Synology NAS storage", "Local backup target and optional uploaded-file storage."),
            ("Offsite backup storage", "Recommended second backup copy: Backblaze B2, AWS S3, Wasabi, or Synology C2."),
            ("Optional firewall/router", "pfSense, OPNsense, MikroTik, or UniFi for stronger network control."),
        ],
        [1.8, 5.2],
    )

    add_heading(doc, "3. Ubuntu Server Setup", 1)
    add_numbered(doc, [
        "Log in to the Ubuntu server with SSH.",
        "Update the server and install Python, PostgreSQL, Nginx, UFW, Fail2ban, Git, and backup tools.",
        "Create a dedicated Linux user named ats.",
        "Place the project in /home/ats/app.",
        "Create a production .env file with SECRET_KEY and DATABASE_URL.",
        "Create a Python virtual environment and install requirements plus gunicorn.",
        "Test Gunicorn locally before exposing the site.",
    ])
    add_code(doc, "sudo apt update && sudo apt upgrade -y")
    add_code(doc, "sudo apt install -y python3 python3-venv python3-pip postgresql postgresql-contrib nginx ufw fail2ban git curl unzip rsync postgresql-client")
    add_code(doc, "sudo adduser ats\nsudo usermod -aG sudo ats\nsu - ats\nmkdir -p /home/ats/app")
    add_code(doc, "cd /home/ats/app\npython3 -m venv .venv\nsource .venv/bin/activate\npip install --upgrade pip\npip install -r requirements.txt\npip install gunicorn")

    add_heading(doc, "4. PostgreSQL Setup", 1)
    add_numbered(doc, [
        "Create the ats database.",
        "Create a non-admin database user named ats_user.",
        "Use a long random password.",
        "Do not expose PostgreSQL port 5432 to the internet.",
    ])
    add_code(doc, "sudo -u postgres psql")
    add_code(doc, "CREATE DATABASE ats;\nCREATE USER ats_user WITH ENCRYPTED PASSWORD 'REPLACE_WITH_LONG_RANDOM_DB_PASSWORD';\nGRANT ALL PRIVILEGES ON DATABASE ats TO ats_user;\n\\q")
    add_code(doc, "sudo -u postgres psql -d ats\nGRANT ALL ON SCHEMA public TO ats_user;\nALTER SCHEMA public OWNER TO ats_user;\n\\q")

    add_heading(doc, "5. Production Environment File", 1)
    add_code(doc, "cd /home/ats/app\nnano .env")
    add_code(doc, "SECRET_KEY=REPLACE_WITH_LONG_RANDOM_SECRET_KEY\nDATABASE_URL=postgresql://ats_user:REPLACE_WITH_LONG_RANDOM_DB_PASSWORD@localhost:5432/ats\nFLASK_ENV=production")

    add_heading(doc, "6. Gunicorn Systemd Service", 1)
    add_code(doc, "sudo nano /etc/systemd/system/ats.service")
    add_code(doc, "[Unit]\nDescription=Applicant Tracking System Flask App\nAfter=network.target postgresql.service\n\n[Service]\nUser=ats\nGroup=www-data\nWorkingDirectory=/home/ats/app\nEnvironmentFile=/home/ats/app/.env\nExecStart=/home/ats/app/.venv/bin/gunicorn -w 4 -b 127.0.0.1:5000 app:app\nRestart=always\nRestartSec=5\n\n[Install]\nWantedBy=multi-user.target")
    add_code(doc, "sudo systemctl daemon-reload\nsudo systemctl enable ats\nsudo systemctl start ats\nsudo systemctl status ats")

    add_heading(doc, "7. Nginx Reverse Proxy", 1)
    p = doc.add_paragraph()
    p.add_run("Reverse proxy meaning: ").bold = True
    p.add_run("Nginx is the public front door on ports 80/443. It forwards requests internally to Flask/Gunicorn on 127.0.0.1:5000. This keeps Flask off the public internet and lets Nginx handle HTTPS, static files, rate limiting, and buffering.")
    add_code(doc, "sudo nano /etc/nginx/sites-available/ats")
    add_code(doc, "server {\n    listen 80;\n    server_name yourdomain.com www.yourdomain.com;\n    client_max_body_size 5M;\n\n    add_header X-Frame-Options \"SAMEORIGIN\" always;\n    add_header X-Content-Type-Options \"nosniff\" always;\n    add_header Referrer-Policy \"strict-origin-when-cross-origin\" always;\n\n    location /static/ {\n        alias /home/ats/app/static/;\n        expires 7d;\n        add_header Cache-Control \"public\";\n    }\n\n    location / {\n        proxy_pass http://127.0.0.1:5000;\n        proxy_http_version 1.1;\n        proxy_set_header Host $host;\n        proxy_set_header X-Real-IP $remote_addr;\n        proxy_set_header X-Forwarded-For $proxy_add_x_forwarded_for;\n        proxy_set_header X-Forwarded-Proto $scheme;\n    }\n}")
    add_code(doc, "sudo ln -s /etc/nginx/sites-available/ats /etc/nginx/sites-enabled/ats\nsudo nginx -t\nsudo systemctl reload nginx")

    add_heading(doc, "8. Firewall And Router", 1)
    add_bullets(doc, [
        "Ubuntu UFW: allow only OpenSSH, 80/tcp, and 443/tcp.",
        "Router port forwarding: forward public TCP 80 and 443 only to the Ubuntu server LAN IP.",
        "Do not forward Flask 5000, PostgreSQL 5432, Synology DSM 5000/5001, SMB 445, rsync 873, or NFS 2049.",
    ])
    add_code(doc, "sudo ufw default deny incoming\nsudo ufw default allow outgoing\nsudo ufw allow OpenSSH\nsudo ufw allow 80/tcp\nsudo ufw allow 443/tcp\nsudo ufw enable\nsudo ufw status verbose")

    add_heading(doc, "9. Hostinger Domain And Cloudflare DNS", 1)
    add_numbered(doc, [
        "In Hostinger hPanel, go to Domains > Get a new domain, buy the domain, enable privacy if available, and enable auto-renewal.",
        "In Cloudflare, click Add a site, enter the domain, and choose Free or Pro.",
        "In Cloudflare DNS, add an A record for @ pointing to your static public IPv4. Set proxy status to Proxied.",
        "Add a CNAME record for www pointing to your root domain. Set proxy status to Proxied.",
        "Cloudflare will show two nameservers.",
        "In Hostinger hPanel, go to Domains > Domain portfolio > Manage > DNS / Nameservers > Change Nameservers.",
        "Enter the two Cloudflare nameservers and save.",
        "Wait for DNS propagation, usually minutes to 24 hours.",
    ])
    add_table(
        doc,
        ["Type", "Name", "Target", "Proxy"],
        [
            ("A", "@", "YOUR_PUBLIC_STATIC_IP", "Proxied"),
            ("CNAME", "www", "yourdomain.com", "Proxied"),
        ],
        [0.8, 1.0, 3.3, 1.2],
    )

    add_heading(doc, "10. HTTPS Certificate", 1)
    add_code(doc, "sudo apt install -y certbot python3-certbot-nginx\nsudo certbot --nginx -d yourdomain.com -d www.yourdomain.com\nsudo certbot renew --dry-run")
    add_bullets(doc, [
        "In Cloudflare SSL/TLS, set mode to Full (strict).",
        "Enable Always Use HTTPS.",
        "Confirm https://yourdomain.com loads before announcing the system.",
    ])

    add_heading(doc, "11. Synology Backup Configuration", 1)
    add_numbered(doc, [
        "In Synology DSM, open Control Panel > Shared Folder > Create > Create Shared Folder.",
        "Create a shared folder named ats_backups.",
        "Enable Recycle Bin. Enable encryption if your Synology supports it.",
        "Open Control Panel > User & Group > Create.",
        "Create a non-admin user named ats_backup.",
        "Give ats_backup Read/Write access only to ats_backups and No Access to other folders.",
        "Open Control Panel > File Services > rsync and enable rsync service.",
        "Open Control Panel > Application Privileges > rsync > Edit and allow ats_backup. Restrict to the Ubuntu server LAN IP if DSM provides that option.",
    ])
    add_code(doc, "nano /home/ats/backup_to_synology.sh")
    add_code(doc, "#!/bin/bash\nset -e\nDATE=$(date +%F-%H%M)\nLOCAL_BACKUP_DIR=\"/var/backups/ats\"\nDB_URL=\"postgresql://ats_user:YOUR_DB_PASSWORD@localhost:5432/ats\"\nSYNOLOGY_IP=\"192.168.1.20\"\nSYNOLOGY_USER=\"ats_backup\"\nSYNOLOGY_TARGET=\"ats_backups\"\nmkdir -p \"$LOCAL_BACKUP_DIR\"\npg_dump \"$DB_URL\" | gzip > \"$LOCAL_BACKUP_DIR/ats-db-$DATE.sql.gz\"\nrsync -avz \"$LOCAL_BACKUP_DIR/\" \"$SYNOLOGY_USER@$SYNOLOGY_IP:/$SYNOLOGY_TARGET/\"\nfind \"$LOCAL_BACKUP_DIR\" -type f -name \"*.gz\" -mtime +7 -delete")
    add_code(doc, "chmod +x /home/ats/backup_to_synology.sh\n/home/ats/backup_to_synology.sh\ncrontab -e")
    add_code(doc, "0 2 * * * /home/ats/backup_to_synology.sh >> /home/ats/backup.log 2>&1")

    add_heading(doc, "12. Uploaded PDFs And Excel Files", 1)
    add_bullets(doc, [
        "Current behavior: applicant CV PDFs are stored in PostgreSQL as applicant.cv_data.",
        "Current behavior: employee profile images are stored in PostgreSQL as employee.profile_data.",
        "Current behavior: Excel exports are generated in memory and downloaded; they are not permanently stored by the app.",
        "Because PDFs are currently inside PostgreSQL, the Synology database backup includes the PDFs.",
        "Optional future improvement: change the app to store uploads on /mnt/ats_uploads and keep only file path, filename, content type, size, and checksum in PostgreSQL.",
    ])

    add_heading(doc, "13. Optional Direct Synology Upload Storage", 1)
    add_numbered(doc, [
        "In DSM, create another shared folder named ats_uploads.",
        "Create a non-admin user named ats_storage.",
        "Give ats_storage Read/Write access only to ats_uploads.",
        "On Ubuntu, mount the Synology share at /mnt/ats_uploads using CIFS.",
        "Only after the mount is reliable, modify the Flask app to write PDFs/profile images to that mount instead of database binary columns.",
    ])
    add_code(doc, "sudo apt install -y cifs-utils\nsudo mkdir -p /mnt/ats_uploads\nsudo nano /etc/ats-synology-credentials")
    add_code(doc, "username=ats_storage\npassword=YOUR_STRONG_SYNOLOGY_PASSWORD")
    add_code(doc, "sudo chmod 600 /etc/ats-synology-credentials\nsudo nano /etc/fstab")
    add_code(doc, "//192.168.1.20/ats_uploads /mnt/ats_uploads cifs credentials=/etc/ats-synology-credentials,uid=ats,gid=www-data,file_mode=0640,dir_mode=0750,iocharset=utf8,nofail,_netdev 0 0\nsudo mount -a")

    add_heading(doc, "14. Security Checklist", 1)
    add_bullets(doc, [
        "Use Cloudflare proxy for @ and www DNS records.",
        "Expose only ports 80 and 443 to the internet.",
        "Keep Synology LAN-only; never port-forward DSM, SMB, rsync, or NFS.",
        "Disable SSH password login after SSH key login is confirmed.",
        "Use Fail2ban.",
        "Use strong unique passwords for PostgreSQL, Synology users, Ubuntu users, and Cloudflare/Hostinger.",
        "Enable 2FA on Hostinger, Cloudflare, Synology, and admin email accounts.",
        "Back up PostgreSQL nightly to Synology and keep at least one offsite backup copy.",
        "Test restores regularly, not only backups.",
        "Review application logs and Nginx logs weekly during launch.",
    ])

    add_heading(doc, "15. Deployment Update Procedure", 1)
    add_code(doc, "cd /home/ats/app\ngit pull\nsource .venv/bin/activate\npip install -r requirements.txt\nsudo systemctl restart ats\nsudo systemctl status ats\njournalctl -u ats -n 100 --no-pager")

    add_heading(doc, "16. Go-Live Verification", 1)
    add_bullets(doc, [
        "https://yourdomain.com loads successfully.",
        "Applicant registration form works.",
        "PDF CV upload works.",
        "Employee login works.",
        "Dashboard loads and can download CVs.",
        "Excel export works.",
        "Cloudflare DNS records are proxied.",
        "External scan shows only ports 80 and 443 open.",
        "PostgreSQL 5432, Flask 5000, and Synology ports are closed externally.",
        "Synology backup script has produced a recent .sql.gz backup.",
    ])

    add_heading(doc, "17. Source References", 1)
    add_bullets(doc, [
        "Hostinger DNS records: https://support.hostinger.com/en/articles/1583249-how-to-manage-dns-records-at-hostinger",
        "Hostinger A records: https://support.hostinger.com/en/articles/4468886-how-to-add-and-remove-a-records-in-hpanel/",
        "Docker on Ubuntu: https://docs.docker.com/engine/install/ubuntu/",
        "Gunicorn deployment: https://docs.gunicorn.org/en/19.9.0/deploy.html",
        "Cloudflare Free plan: https://www.cloudflare.com/plans/free/",
        "Cloudflare Universal SSL: https://developers.cloudflare.com/ssl/edge-certificates/universal-ssl/",
        "Synology shared folders: https://kb.synology.com/en-af/PAS/help/PAS/AdminCenter/file_share_create",
        "Synology rsync: https://kb.synology.com/en-us/DSM/help/DSM/AdminCenter/file_rsync?version=7",
    ])

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("ARTS Deployment Procedure")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUT)


if __name__ == "__main__":
    build()

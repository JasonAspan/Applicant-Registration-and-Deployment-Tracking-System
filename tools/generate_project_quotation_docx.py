from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path("Project_Quotation_ATS_Server_Procurement.docx")

BLACK = RGBColor(0, 0, 0)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY = RGBColor(90, 90, 90)
LIGHT_GRAY = "F2F4F7"
WHITE = RGBColor(255, 255, 255)

USD_TO_PHP = 61.38


def php(amount_usd: float) -> float:
    return round(amount_usd * USD_TO_PHP, 2)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcw = tc_pr.find(qn("w:tcW"))
    if tcw is None:
        tcw = OxmlElement("w:tcW")
        tc_pr.append(tcw)
    tcw.set(qn("w:w"), str(width_dxa))
    tcw.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_borders(table, color="BFC7D1", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run(run, size=11, color=BLACK, bold=False, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic


def add_p(doc, text="", style=None, size=11, color=BLACK, bold=False, italic=False, after=6, before=0, align=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run(run, size=size, color=color, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12 if level == 2 else 8)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6 if level == 2 else 4)
    run = p.add_run(text)
    set_run(run, size=16 if level == 1 else 13 if level == 2 else 12, color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_key_value(doc, rows):
    widths = [1800, 7560]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = ""
        cells[1].text = ""
        p0 = cells[0].paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        r0 = p0.add_run(label)
        set_run(r0, bold=True)
        p1 = cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        r1 = p1.add_run(value)
        set_run(r1)
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_table_geometry(table, widths)
    set_cell_margins(table)
    set_borders(table)
    return table


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        set_cell_shading(hdr[i], LIGHT_GRAY)
        p = hdr[i].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        set_run(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=10.5)
            if i >= len(row) - 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_table_geometry(table, widths)
    set_cell_margins(table)
    set_borders(table)
    return table


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for idx, size in [(1, 16), (2, 13), (3, 12)]:
        style = styles[f"Heading {idx}"]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLUE if idx < 3 else DARK_BLUE
        style.paragraph_format.space_before = Pt(16 if idx == 1 else 12 if idx == 2 else 8)
        style.paragraph_format.space_after = Pt(8 if idx == 1 else 6 if idx == 2 else 4)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("Project Quotation | Applicant Tracking System")
    set_run(r, size=9, color=GRAY)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Prepared for review and approval")
    set_run(r, size=9, color=GRAY)


def build():
    doc = Document()
    configure_doc(doc)

    add_p(doc, "PROJECT QUOTATION", size=24, color=BLACK, bold=True, after=2)
    add_p(doc, "Applicant Tracking System - Server Procurement and Deployment", size=14, color=GRAY, after=14)
    add_key_value(
        doc,
        [
            ("Prepared for", "[Client / Company Name]"),
            ("Prepared by", "[Supplier / Service Provider Name]"),
            ("Quotation date", "May 25, 2026"),
            ("Validity", "15 calendar days from quotation date"),
            ("Currency", "PHP, with USD cloud pricing converted at PHP 61.38 per USD"),
            ("Document status", "For budget approval and server purchase decision"),
        ],
    )

    add_heading(doc, "Executive Summary", 1)
    add_p(
        doc,
        "This quotation covers the recommended server package for hosting the Applicant Tracking System (ATS), a Flask-based web application with employee authentication, RBAC administration, applicant records, PDF CV uploads, PostgreSQL support, and employee dashboard workflows.",
    )
    add_p(
        doc,
        "The recommended starter production setup is a cloud VPS with PostgreSQL, automated backups, SSL, deployment hardening, and a separate small staging server. This is sized for an initial small-to-medium internal deployment and can be upgraded without changing the application architecture.",
    )

    add_heading(doc, "Recommended Server Architecture", 1)
    add_table(
        doc,
        ["Component", "Recommended Specification", "Purpose"],
        [
            ("Production application server", "2 vCPU, 4 GB RAM, 80 GB SSD, Ubuntu LTS, public IPv4", "Runs Flask/Gunicorn, Nginx reverse proxy, static assets, upload handling, and application services."),
            ("Database", "PostgreSQL 16, managed database preferred; self-managed PostgreSQL acceptable for lower monthly cost", "Stores employees, roles, permissions, applicants, status history, and uploaded document metadata."),
            ("Backup layer", "Daily cloud backups plus periodic manual snapshots", "Recovery from server failure, accidental deletion, or deployment rollback."),
            ("Staging server", "1 vCPU, 1 GB RAM, 25 GB SSD", "Testing updates before production release."),
            ("Security", "HTTPS, firewall, environment secrets, non-root deployment user, restricted database access", "Protects employee accounts, applicant records, and uploaded CVs."),
        ],
        [1900, 3700, 3760],
    )

    add_heading(doc, "Monthly Cloud Server Quotation", 1)
    monthly_rows = [
        ("Production VPS", "DigitalOcean Basic Droplet, 4 GiB RAM, 2 vCPU, 80 GiB SSD, 4,000 GiB transfer", "USD 24.00", f"PHP {php(24):,.2f}"),
        ("Daily backup", "DigitalOcean backup estimate at 30% of production VPS cost", "USD 7.20", f"PHP {php(7.2):,.2f}"),
        ("Manual snapshot reserve", "80 GB snapshot allowance at USD 0.06/GB-month", "USD 4.80", f"PHP {php(4.8):,.2f}"),
        ("Staging VPS", "DigitalOcean Basic Droplet, 1 GiB RAM, 1 vCPU, 25 GiB SSD, 1,000 GiB transfer", "USD 6.00", f"PHP {php(6):,.2f}"),
        ("Managed PostgreSQL", "Starter managed PostgreSQL allowance; can be deferred if PostgreSQL is self-hosted on the production VPS", "USD 15.00", f"PHP {php(15):,.2f}"),
    ]
    add_table(doc, ["Item", "Description", "Monthly USD", "Monthly PHP"], monthly_rows, [1700, 4300, 1500, 1860])
    monthly_total = 24 + 7.2 + 4.8 + 6 + 15
    add_p(doc, f"Estimated monthly operating cost: USD {monthly_total:,.2f} / PHP {php(monthly_total):,.2f}.", bold=True, after=2)
    add_p(doc, f"Estimated 12-month operating cost: USD {monthly_total * 12:,.2f} / PHP {php(monthly_total * 12):,.2f}.", bold=True)

    add_heading(doc, "One-Time Implementation Services", 1)
    service_rows = [
        ("Server provisioning", "Create VPS resources, configure Ubuntu, users, firewall, package updates, timezone, and baseline security.", "PHP 12,000.00"),
        ("Application deployment", "Deploy ATS codebase with Python environment, Gunicorn service, Nginx reverse proxy, production environment variables, and static/upload configuration.", "PHP 18,000.00"),
        ("Database setup", "Provision PostgreSQL, initialize schema, configure connection string, seed initial RBAC roles/permissions, and verify login/admin flow.", "PHP 10,000.00"),
        ("SSL and domain setup", "Configure domain records, HTTPS certificate, renewal path, and secure headers.", "PHP 6,000.00"),
        ("Backup and recovery setup", "Configure provider backups, snapshot schedule, database export routine, and restore checklist.", "PHP 8,000.00"),
        ("Production verification", "Smoke test applicant submission, employee login, dashboard, PDF CV upload/download, batch export, admin panel, and RBAC restrictions.", "PHP 9,000.00"),
    ]
    add_table(doc, ["Service", "Scope", "Amount"], service_rows, [2200, 5160, 2000])
    one_time_total = 63000
    add_p(doc, f"One-time implementation services total: PHP {one_time_total:,.2f}.", bold=True)

    add_heading(doc, "Quotation Summary", 1)
    summary_rows = [
        ("One-time implementation services", "PHP", f"{one_time_total:,.2f}"),
        ("First month cloud server operating cost", "PHP", f"{php(monthly_total):,.2f}"),
        ("Estimated first payment", "PHP", f"{one_time_total + php(monthly_total):,.2f}"),
        ("Estimated annual cloud operating cost", "PHP", f"{php(monthly_total * 12):,.2f}"),
    ]
    add_table(doc, ["Cost Category", "Currency", "Amount"], summary_rows, [5000, 1600, 2760])

    add_heading(doc, "Optional Cost Reduction", 1)
    add_p(
        doc,
        "If the initial budget is tight, the managed PostgreSQL item can be deferred and PostgreSQL can run on the production VPS using Docker. This lowers the monthly estimate by USD 15.00, reducing the monthly operating cost to USD 42.00 / PHP 2,578.00. The tradeoff is that database patching, backups, and recovery become the project team's operational responsibility.",
    )

    add_heading(doc, "Scope Included", 1)
    included = [
        "Production server provisioning and hardening",
        "PostgreSQL configuration for the ATS",
        "Application deployment and service configuration",
        "Nginx reverse proxy and HTTPS setup",
        "Backup and restore procedure setup",
        "Production smoke testing for applicant, employee, dashboard, CV upload/download, batch export, admin panel, and RBAC flows",
        "Basic handover notes for server access, restart commands, and backup checks",
    ]
    for item in included:
        add_p(doc, item, style="List Bullet", after=4)

    add_heading(doc, "Assumptions", 1)
    assumptions = [
        "The ATS will initially support a small-to-medium internal recruitment workflow, not a high-traffic public job board.",
        "PDF CV uploads are limited by application configuration and should be monitored because uploaded files affect database and backup growth.",
        "The client will provide a domain name, official sender email details if email notifications are later added, and cloud account billing access.",
        "Prices are cloud infrastructure estimates and may change before purchase. Final billing is controlled by the cloud provider.",
        "This quotation does not include paid third-party SMS, email, antivirus scanning, SSO, 2FA service subscriptions, or penetration testing unless separately approved.",
    ]
    for item in assumptions:
        add_p(doc, item, style="List Bullet", after=4)

    add_heading(doc, "Payment Terms", 1)
    add_key_value(
        doc,
        [
            ("Payment schedule", "50% upon approval, 50% upon production handover."),
            ("Cloud billing", "Paid directly by the client to the cloud provider unless otherwise agreed."),
            ("Lead time", "2 to 4 working days after cloud account access, domain access, and initial payment are available."),
            ("Warranty", "7 calendar days post-handover for deployment defects related to the included scope."),
            ("Change requests", "New features, UI changes, migration cleanup, or integrations are quoted separately."),
        ],
    )

    add_heading(doc, "Reference Pricing Sources", 1)
    sources = [
        "DigitalOcean Droplet pricing page, accessed May 25, 2026: https://www.digitalocean.com/pricing/droplets",
        "DigitalOcean Managed Databases pricing page, accessed May 25, 2026: https://www.digitalocean.com/pricing/managed-databases",
        "USD/PHP reference rate used for this estimate: PHP 61.38 per USD, accessed May 25, 2026: https://exchangerate.guru/usd/php/1/",
        "Amazon Lightsail billing reference for comparison, accessed May 25, 2026: https://docs.aws.amazon.com/lightsail/latest/userguide/amazon-lightsail-frequently-asked-questions-faq-billing-and-account-management.html",
    ]
    for item in sources:
        add_p(doc, item, size=9.5, color=GRAY, after=4)

    add_heading(doc, "Approval", 1)
    add_table(
        doc,
        ["Role", "Name and Signature", "Date"],
        [
            ("Client authorized representative", "", ""),
            ("Service provider authorized representative", "", ""),
        ],
        [2600, 4760, 2000],
    )

    doc.core_properties.title = "Project Quotation - Applicant Tracking System Server Procurement"
    doc.core_properties.subject = "Server procurement and deployment quotation"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generated for ATS server procurement review."
    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT.resolve())
